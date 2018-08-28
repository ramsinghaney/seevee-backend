
# coding: utf-8

# In[8]:


import os
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from docx import Document
import string
import re
import csv
import numpy as np
bar_value=0


stop_words = set(stopwords.words('english'))
punc=set(string.punctuation)
parent_list=['a2billing', 'adk', 'ado.net', 'ado.net entity framework', 'ajax', 'amqp', 'ansi c', 'api development', 'api documentation', 'arm', 'asp', 'asp.net', 'asp.net mvc', 'atl', 'atm implementation', 'ablecommerce', 'abstract window toolkit (awt)', 'actian', 'actionscript', 'actionscript 3', 'activex', 'activex data objects (ado)', 'ada', 'adaptive algorithms', 'adobe air', 'adobe dreamweaver', 'adobe', 'adobe flash', 'adobe flex', 'advanced business application programming (abap)', 'agile software development', 'akka', 'alfresco development', 'alfresco user', 'algorithm development', 'alpha', 'amazon appstore', 'amazon ec2', 'amazon mws', 'amazon relational database service', 'amazon s3', 'amazon web services', 'android', 'android app development', 'android sdk', 'angularjs', 'apache administration', 'apache ant', 'apache avro', 'apache cxf', 'apache camel', 'apache cassandra', 'apache click', 'apache cocoon', 'apache cordova', 'apache flume', 'apache hive', 'apache jakarta poi', 'apache kafka', 'apache mahout', 'apache maven', 'apache ofbiz', 'apache shirol', 'apache solr', 'apache spark', 'apache struts', 'apache subversion (svn)', 'apache thrift', 'apache tiles', 'apache tomcat', 'apollo', 'app store', 'app usability analysis', 'appfuse', 'appcelerator titanium', 'appian bpm suite', 'apple uikit framework', 'apple webobjects', 'apple xcode', 'applescript', 'application lifecycle management', 'application programming', 'application server', 'arcgis', 'arcscene', 'arduino', 'artificial intelligence', 'artificial neural networks', 'artioscad', 'artisteer', 'aspdotnetstorefront', 'assembla', 'assembly language', 'asynchronous i/o', 'atlassian confluence', 'atlassian jira', 'atmel avr', 'atom', 'auctiva', 'augmented reality', 'authorize.net development', 'autohotkey', 'autodys accelicad', 'autoit', 'automated testing', 'automation', 'awk', 'axapta', 'axiis', 'axiom microstation productivity toolkit', 'axiom productivity tools', 'axure rp', 'amember', 'aweber', '.net compact framework', '.net framework', '.net remoting', 'bgl simple fund', 'birt', 'bpcs', 'backbone.js', 'bada', 'basecamp', 'bash shell scripting', 'beos', 'behavior driven development (bdd)', 'big data', 'bigcommerce', 'bitrock installbuilder', 'bitcoin', 'bitrix', 'bitrix intranet', 'biztalk server', 'black box testing', 'blackberry app development', 'blackberry jde', 'blazeds', 'blitz basic', 'bluetooth', 'boonex dolphin', 'boost', 'bootstrap', 'borland c++ builder', 'borland silktest', 'box.net development', 'box2d', 'buddypress', 'bugzilla', 'business activity monitoring', 'business it alignment', 'business intelligence', 'business process modeling', 'business process reengineering', 'bbpress', 'c', 'c#', 'c++', 'cgi', 'cms development', 'cnc programming', 'cobol', 'corba', 'cpanel', 'crm', 'cs-cart', 'css', 'css3', 'cuda', 'curl', 'cvs', 'cache management', 'cairngorm', 'cakephp', 'calculus', 'capistrano', 'capture nx2', 'carbide.c++', 'caspio administration', 'caspio programming', 'central desktop development', 'centreon', 'certified information systems security professional (cissp)', 'chrome extension', 'chrome os', 'civicrm', 'clearquest', 'clickbank', 'clojure', 'cloud security framework', 'cloudforge', 'codesys', 'cocoa', 'cocoa touch', 'cocos2d', 'code refactoring', 'codeigniter', 'codewarrior', 'coffeescript', 'cognos', 'coldfusion', 'collabnet teamforge', 'collaborative filtering', 'comet', 'communications', 'compiler', 'component object model (microsoft com)', 'computational linguistics', 'computer engineering', 'computer vision', 'concept software inpage', 'concrete5 cms', 'constant contact', 'contao cms', 'continuous integration', 'core java', 'corel ventura', 'corona', 'cosmos os', 'cryptography', 'cubecart', 'cucumber', 'custom cms', 'customer information control system (cics)', 'd programming language', 'dart', 'dbms', 'dos', 'davinci resolve', 'data backup', 'data cleansing', 'data ingestion', 'data logistics', 'data modeling', 'data science', 'data scraping', 'data structures', 'data visualization', 'data warehousing', 'datalife engine', 'database caching', 'database cataloguing', 'database design', 'database modeling', 'database programming', 'database testing', 'delphi', 'demandware', 'desktop applications', 'devexpress', 'devops', 'device driver development', 'digital access pass', 'digital engineering', 'digital mapping', 'digital signal processing', 'directshow', 'directx', 'distributed computing', 'django', 'docbook', 'doctrine orm', 'document object model', 'dojo toolkit', 'domain migration', 'dotnetnuke', 'dropbox api', 'drupal', 'dwolla api', 'd3.js', 'dbase administration', 'dbase programming', 'ecmascript', 'edge', 'erdas imagine', 'eclipse', 'ecommerce platform development', 'elasticsearch', 'elastix', 'electronic data interchange (edi)', 'electronic workbench', 'elgg', 'embedded c', 'embedded linux', 'embedded systems', 'ember.js', 'enterprise javabeans (ejb)', 'entity framework', 'erlang', 'erwin', 'eucalyptus cloud', 'excel vba', 'express scribe', 'expression engine', 'ext js', 'e-learning', 'ebay api', 'ez publish', 'f#', 'fbml', 'ffmpeg', 'ftp', 'facebook development', 'facebook games development', 'facebook javascript (fbjs)', 'field-map', 'filemaker', 'fire os development', 'firebird', 'firefox plugin development', 'flash 3d', 'flask', 'flowcharts', 'foreign exchange trading', 'form-z', 'fortran', 'forum development', 'foursquare development', 'foxpro programming', 'freemarker', 'freeswitch', 'frontend development', 'full-text search engines', 'functional testing', '', 'gimp', 'glsl', 'gnu octave', 'gps development', 'gtk+', 'game development', 'game programming', 'game testing', 'gamesalad creator', 'gamification', 'gearman', 'genetic algorithms', 'geographic information system (gis)', 'geolocation', 'geometry', 'geospatial', 'getresponse', 'ggplot2', 'git', 'github', 'glassfish', 'go', 'gotomypc', 'golang', 'goldmine', 'google adsense api', 'google adwords development', 'google analytics api', 'google app engine', 'google app engine api', 'google apps api', 'google calendar development', 'google docs api', 'google gadgets', 'google gadgets api', 'google glass', 'google glass sdk', 'google map maker', 'google maps api', 'google reader api', 'google shopping', 'google sites api', 'google sites administration', 'google swiffy', 'google web toolkit', 'google+ development', 'google+ marketing', 'gradle', 'grails', 'graph databases', 'graphics programming', 'gravity forms', 'groovy', 'gruntjs', 'haml', 'hbase', 'hp cloud', 'hp network management center (hpnmc)', 'hp quicktest professional (hpqtp)', 'hp-ux administration', 'html', 'html5', 'haxe', 'hadoop', 'haskell', 'heroku', 'hibernate', 'highcharts', 'home automation', 'ibm as/400 control language', 'ibm db2 administration', 'ibm db2 programming', 'ibm lotus domino', 'ibm lotus notes traveler', 'ibm lotus symphony', 'ibm powerpc programming', 'ibm rational rose', 'ibm smartcloud', 'ibm system p', 'ibm system x', 'ibm tivoli framework', 'ibm watson', 'ims', 'imacros', 'it management', 'it service management', 'icefaces', 'image processing', 'in-game advertising', 'indexing', 'informatica', 'information architecture', 'information design', 'infragistics', 'infusionsoft development', 'inno setup', 'instagram api', 'installshield', 'installer development', 'instrumentation', 'integrated circuits', 'intellij idea', 'interactive voice response', 'internet information services (iis)', 'internet security', 'interspire', 'intranet architecture', 'intranet implementation', 'ionic framework', 'issue tracking systems', 'ios development', 'ipad app development', 'iphone app development', 'ireport', 'itextsharp', 'j2ee', 'j2me', 'j2se', 'jaxb', 'jbpm', 'jboss', 'jboss seam', 'jcl', 'jdbc', 'jdeveloper', 'jmeter', 'jncia-junos', 'jndi', 'jpa', 'jquery mobile', 'json', 'jsp', 'jstl', 'junit', 'jasperreports', 'java', 'java ee', 'java me', 'java remote method invocation (java rmi)', 'java servlets development', 'javafx', 'javascript', 'javaserver faces (jsf)', 'jenkins', 'jetpack', 'jimdo', 'jingle program production', 'jinja2', 'jomsocial development', 'joomla', 'joomla fabrik', 'joomla migration', 'jsharp', 'jquery', 'kvm', 'kajabi', 'kannada', 'kendo ui', 'kentico cms', 'kernel', 'keyboarding', 'kindle app development', 'kindle fire', 'knockoutjs', 'kohana', 'korn shell', 'lamp administration', 'less', 'linq', 'linq to entities', 'linq to sql', 'labview', 'landing pages', 'laravel framework', 'lean consulting', 'lemonstand', 'liferay', 'limejs', 'limesurvey', 'linear programming', 'linkedin development', 'liquidplanner', 'lisp', 'lithium framework', 'load testing', 'loadrunner', 'log4j', 'logixml', 'lotus approach', 'lotus notes', 'lua', 'lucene search', 'libgdx', 'libcurl', 'midi', 'modx', 'mql 4', 'mvc framework', 'mxml', 'mac os app development', 'mac osx administration', 'machine design', 'machine learning', 'magento', 'magic tricks', 'mailchimp', 'make build script', 'malware', 'mambo', 'mantis', 'manual test execution', 'mapreduce', 'mapinfo', 'mcafee virusscan', 'mechanical turk api', 'medical imaging', 'medical informatics', 'memcached', 'merchantrun', 'merchantrun globallink', 'mercurial', 'metatrader 4 (mt4)', 'meteor', 'microstrategy', 'microcontroller programming', 'microsoft access administration', 'microsoft access programming', 'microsoft dynamics administration', 'microsoft dynamics crm', 'microsoft dynamics development', 'microsoft dynamics erp', 'microsoft dynamics gp', 'microsoft entity framework', 'microsoft foundation classes (mfc)', 'microsoft infopath', 'microsoft kinect development', 'microsoft message queue server (mmsq)', 'microsoft office sharepoint server', 'microsoft outlook development', 'microsoft sql ssas', 'microsoft sql ssrs', 'microsoft sql server notification services', 'microsoft sql server programming', 'microsoft sharepoint administration', 'microsoft sharepoint designer', 'microsoft sharepoint development', 'microsoft silverlight', 'microsoft visual c++', 'microsoft visual studio', 'microsoft visual studio lightswitch', 'microsoft windows azure', 'microsoft windows media connect', 'microsoft windows phone 7 app development', 'microsoft windows template library', 'microsoft windows workflow foundation', 'minecraft', 'mobile advertising', 'mobile app development', 'mobile app testing', 'mobile development framework', 'mobile programming', 'mocha', 'mockito', 'model view viewmodel (mvvm)', 'mongodb', 'mono', 'moodle', 'mootools', 'movabletype', 'mozenda scraper', 'multi-touch hardware programming', 'multithreaded programming', 'music', 'mysql administration', 'mysql programming', 'nhibernate', 'natural language processing', 'navigation system design', 'navigation system implementation', 'neo4j', 'netbeans', 'netsuite development', 'network pentesting', 'network programming', 'nginx', 'ning development', 'nosql', 'node.js', 'non-disclosure agreements', 'nopcommerce', 'n2cms', 'oauth', 'ocr algorithms', 'ocr tesseract', 'ocaml', 'odbc', 'ogre', 'olap', 'oops', 'orm', 'osgi', 'object oriented design', 'object oriented php', 'object oriented programming (oop)', 'object pascal', 'objective-c', 'objective-j', 'omnigraffle', 'openbsd', 'openbravo pos', 'opencv', 'opencart', 'openemm', 'openerp administration', 'openerp development', 'opengl', 'opengl es', 'openlayers', 'openstack', 'opentok development', 'openvbx', 'openvms', 'openwrt', 'openx', 'optimizepress', 'orace obiee plus', 'oracle apex', 'oracle atg web commerce', 'oracle agile', 'oracle application framework', 'oracle application server', 'oracle brm', 'oracle crm on demand', 'oracle complex events processing', 'oracle data guard', 'oracle database', 'oracle demantra', 'oracle e-business suite', 'oracle endeca', 'oracle enterprise service bus', 'oracle forms', 'oracle fusion applications', 'oracle fusion middleware', 'oracle hyperion planning', 'oracle java ee', 'oracle plsql', 'oracle patching', 'oracle performance tuning', 'oracle policy automation', 'oracle primavera', 'oracle programming', 'oracle rightnow', 'oracle soa suite', 'oracle sun ray', 'oracle taleo', 'oracle team productivity center', 'oracle transportation management', 'oracle unified method', 'oracle universal content management', 'oracle upgrade', 'oracle user productivity kit', 'oracle weblogic', 'orangecrm', 'orchard cms', 'oscommerce', 'oovoo development', 'opensuse', 'pdf conversion', 'php', 'php-nuke', 'plc & scada', 'plc programming', 'pos terminal development', 'prado php framework', 'psd to html', 'psd to mailchimp', 'psd to wordpress', 'psd to xhtml', 'psd2cms', 'pspice', 'palm', 'palm app development', 'palm webos application development', 'papercraft', 'parallels virtual desktop', 'parse mobile app platform', 'pascal', 'paypal development', 'payment gateway integration', 'payment processing', 'paypal integration', 'pentaho', 'perforce', 'performance testing', 'performance tuning', 'perl', 'perl catalyst', 'perl mojolicious', 'perldancer', 'pervasive software', 'phonegap', 'phpbb', 'physics', 'pig', 'platform migration', 'play framework', 'pligg', 'plivo', 'polymer clay sculpting', 'portlets', 'postscript', 'postgresql administration', 'postgresql programming', 'power builder', 'prestashop', 'primefaces', 'private clouds', 'protools', 'product management', 'program management', 'project scheduling', 'prolog', 'prototype javascript framework', 'puppet administration', 'pyqt', 'pylons', 'python', 'python numpy', 'python scipy', 'phpfox', 'phpmyadmin', 'phpmydirectory', 'qa engineering', 'qa management', 'qt', 'quartz scheduler', 'r-hadoop', 'radius', 'rest', 'rets', 'rss', 'rspec', 'rtl', 'rtml', 'rackspace', 'radiant cms', 'raphael js', 'rapid miner', 'rapid prototyping', 'raspberry pi', 'rational unified process (rup)', 'razor template engine', 'real estate idx', 'real time stream processing', 'realbasic', 'recommender systems', 'red hat certified engineer (rhce)', 'red5', 'redis', 'redmine', 'refinery cms', 'regular expressions', 'relational databases', 'remote sensing', 'requirement management', 'requirements analysis', 'resin', 'resource description framework (rdf)', 'reverse engineering', 'rhodes framework', 'robot framework', 'robotics', 'ruby', 'ruby on rails', 'russian language', 'sap abap', 'sap bsp', 'sap business objects', 'sap crm', 'sap crystal reports', 'sap erp', 'sap logistics execution', 'sap sd', 'sas', 'scada', 'scorm', 'sdlx', 'sip', 'smpp', 'sms', 'sms gateway', 'smtp', 'soap', 'sqa', 'sql', 'sql azure', 'sql clr', 'sql programming', 'sql server integration services (ssis)', 'sqlite administration', 'sqlite programming', 'sqr', 'ssh', 'ssi', 'ssl', 'swt', 'saas', 'sage erp accpac', 'salesforce apex', 'salesforce app development', 'salesforce.com', 'sass', 'scala', 'scenario planning', 'scheme', 'scientific computation', 'scrapebox', 'scrapy', 'scripting', 'scripts & utilities', 'scrum', 'section 508 compliance', 'selenium', 'selenium webdriver', 'sencha gxt', 'sencha touch', 'sendmail', 'sentiment analysis', 'serial port interfacing', 'serialization', 'shiva3d', 'shopify', 'shopify templates', 'silex framework', 'silverstripe', 'simple directmedia layer', 'sinatra framework', 'sitebuildit', 'sitecore', 'skadate', 'skype development', 'slackware linux', 'smartfoxserver', 'smarty', 'snort', 'social networking development', 'socialengine', 'socket programming', 'software configuration management', 'software debugging', 'software documentation', 'software licensing', 'software qa testing', 'software testing', 'spamassassin', 'sparx systems enterprise architect', 'sphinx', 'spine js', 'spree', 'spring framework', 'spring security', 'squarespace', 'stakeholder management', 'standard template library (stl)', 'stenography', 'stream processing', 'stripe', 'subversion', 'sugarcrm development', 'surveymonkey', 'swift', 'swing', 'symbian development', 'symfony', 'system analysis', 'system automation', 'systems development', 'soapui', 'tcp/ip', 'typo3', 'talend open studio', 'tastypie', 'taxonomy', 'teaching algebra', 'teaching mathematics', 'team foundation server', 'telerik', 'telerik sitefinity cms', 'templates', 'tesseract', 'test automation', 'test case design', 'test driven development', 'testcomplete', 'testlink', 'testing framework', 'textpattern', 'thai', 'tizen', 'tk', 'tornado', 'tortoise svn', 'transact-sql', 'tropo', 'twig', 'twilio api', 'twitter api', 'twitter bootstrap', 'uml', 'ubercart', 'umbraco', 'unify', 'unit testing', 'unity', 'unix', 'unix shell', 'unreal engine', 'unrealscript', 'usability testing', 'user acceptance testing', 'vb.net', 'vba', 'vbscript', 'vbulletin', 'vhdl', 'vicidial', 'voip software', 'vtk', 'vaadin framework', 'version control', 'vertica', 'video streaming', 'virtual currency', 'virtual machine', 'virtual storage access method (vsam)', 'virtuemart', 'virtuoso', 'visual basic', 'visual dataflex', 'visual foxpro', 'visualforce', 'voicexml', 'volusion', 'vtiger adminstration', 'vtiger development', 'wamp', 'whmcs development', 'web crawler', 'web crawling', 'web hosting', 'web programming', 'web scraping', 'web services', 'web services development', 'web testing', 'webapp pentesting', 'webgl', 'webisode production', 'website development', 'weebly', 'weka', 'wix', 'wicket', 'win32 app development', 'windev', 'windev mobile', 'windows 8 app development', 'windows app development', 'windows communication foundation (wcf)', 'windows forms development', 'windows mobile', 'windows phone', 'windows presentation foundation (wpf)', 'winsock', 'woocommerce', 'wordpress', 'wordpress e-commerce', 'wordpress plugin', 'worldspan', 'wowza media server', 'wrap advertising', 'wxwidgets', 'x-cart', 'x86 assembly language', 'xaml', 'xhtml', 'xml', 'xml web services', 'xml-rpc', 'xmpp', 'xpath', 'xquery', 'xsd', 'xsl', 'xslt', 'xul', 'xamarin', 'xara xtreme', 'xbox', 'xenforo', 'yui library', 'yahoo developer skills', 'yahoo! advertising solutions', 'yahoo! query language', 'yahoo! store', 'yii', 'yola', 'youtube development', 'zk', 'zapier', 'zen cart', 'zend framework', 'zend studio', 'zendesk api development', 'zoho creator', 'zoomla', 'zope', 'zurb foundation']
secondary_list=['c/c++', 'verilog', 'x86', 'java proxy', 'jscript', 'cake', 'cms', 'r', '.net', 'c#.net', 'agile', 'brew', 's-60', 'mysql', 'sql sever', 'oracle', 'html/css', 'qml', 'scipting', 'script', 'shell', 'visual studio', 'windows', 'win32', 'symbian', 'mac', 'linux', 'api/gui', 'cloud', 'bugzila', 'struts', 'dom', 'dhtml', 'dreamweaver', 'aptana', 'firebug', 'spring', 'tomcat', 'dojo', 'basic', 'dbase', 'rbase', 'ingres', 'knse', 'assembly', 'sml', 'apple--visual', 'matlab', 'matlab', 'latex', 'data mining', 'rdbms', 'ms access', 'ms-dos', 't-sql', 'ssis', 'ssrs', 'ssas', 'sharepoint', 'rational rose', 'ant', 'aws', 'docker', 'spark', 'svn', 'emacs', 'jbuilder', 'junit testing', 'smalltalk', 'web designing', 'computer graphics', 'srilm', 'wxpython', 'sun', 'solaris', 'mips', 'spim', 'simulator', 'nltk', 'xilinx-ise', 'auto-cad', 'hdl', 'adobe photoshop', 'visual c++', 'datalog', 'p spice', 'sivaco', 'ms office', 'sap']
def listMatch(skill_list,job_description,n):
    c=0
    for w in skill_list:
        if w in job_description:
            c+=1
    
    try:
        c=c/len(job_description)*100
    except:
        c=0
    
    if c>=n:
        return True
    else:
        return None

def shortenText(text):
    final_words=[]
    word_tokens=word_tokenize(text)
    filtered_sentence = [w.lower() for w in word_tokens if not w.lower() in stop_words if not w in punc]



#     for w in filtered_sentence:
#         if w not in final_words:
#             final_words.append(w)
            
#     return final_words
    return filtered_sentence
    
    
def shortenDoc(path):
    document=Document(path)
    final_words=[]
    word_tokens=[]
    table_word_tokens=[]
    
    
    for p in document.paragraphs: 
        word_tokens += word_tokenize(p.text)

    filtered_sentence = [w.lower() for w in word_tokens if not w.lower() in stop_words if not w in punc]
    
    try:
        tables = document.tables
        for table in tables:
            for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            table_word_tokens += word_tokenize(paragraph.text)

        filtered_sentence += [w.lower() for w in table_word_tokens if not w.lower() in stop_words if not w in punc]
    except:
        i=1



#     for w in filtered_sentence:
#         if w not in final_words:
#             final_words.append(w)
            
#     return final_words
    return filtered_sentence

def experience(final_words):
    total_exp=0
    for i in range(0, len(final_words)):
        if final_words[i]=="years":
            if final_words[i-1]=="plus" or final_words[i-1]=="consecutive":
                word= final_words[i-2]
            else:
                word= final_words[i-1]
            try:
                to_int = int(word)
                total_exp+=to_int
            except:
                if '+' in word:
                    strip_word = word.strip('+')
                    try:
                        to_int = int(strip_word)
                        total_exp+=to_int
                    except:
                        total_exp+=0
            return total_exp
    return 0

    
        



def primarySkillMatch(final_words):
    primary_skills_present = []
    primary_skills_present_bi = []
    primary_skills_present_tri = []
    primary_skills_present_tetra = []
    i=0
    
    for w in final_words:
        
        first_word=w
        
#         third_word=final_words[i+2]
        
        
        if w in parent_list and w not in primary_skills_present:
            primary_skills_present.append(w)
            
        if i <=(len(final_words)-2):
            second_word=final_words[i+1]
            combined_word=first_word+" "+second_word
            if combined_word in parent_list and combined_word not in primary_skills_present_bi:
                primary_skills_present_bi.append(combined_word)
            
            
        if i <=(len(final_words)-3):
            second_word=final_words[i+1]
            third_word=final_words[i+2]
            combined_word=first_word+" "+second_word+" "+third_word
            if combined_word in parent_list and combined_word not in primary_skills_present_tri:
                primary_skills_present_tri.append(combined_word)
                
        if i <=(len(final_words)-4):
            second_word=final_words[i+1]
            third_word=final_words[i+2]
            fourth_word=final_words[i+3]
            combined_word=first_word+" "+second_word+" "+third_word+" "+fourth_word
            if combined_word in parent_list and combined_word not in primary_skills_present_tetra:
                primary_skills_present_tetra.append(combined_word)

        i+=1
        
        
    all_skills=primary_skills_present+primary_skills_present_bi+primary_skills_present_tri+primary_skills_present_tetra
            
        
    return all_skills
    
def secSkillMatch(final_words):    
    
    primary_skills_present = []
    primary_skills_present_bi = []
    primary_skills_present_tri = []
    primary_skills_present_tetra=[]
    i=0
    
    for w in final_words:
        
        first_word=w
        
#         third_word=final_words[i+2]
        
        
        if w in secondary_list and w not in primary_skills_present:
            primary_skills_present.append(w)
            
        if i <=(len(final_words)-2):
            second_word=final_words[i+1]
            combined_word=first_word+" "+second_word
            if combined_word in secondary_list and combined_word not in primary_skills_present_bi:
                primary_skills_present_bi.append(combined_word)
            
            
        if i <=(len(final_words)-3):
            second_word=final_words[i+1]
            third_word=final_words[i+2]
            combined_word=first_word+" "+second_word+" "+third_word
            if combined_word in secondary_list and combined_word not in primary_skills_present_tri:
                primary_skills_present_tri.append(combined_word)
                
        if i <=(len(final_words)-4):
            second_word=final_words[i+1]
            third_word=final_words[i+2]
            fourth_word=final_words[i+3]
            combined_word=first_word+" "+second_word+" "+third_word+" "+fourth_word
            if combined_word in parent_list and combined_word not in primary_skills_present_tetra:
                primary_skills_present_tetra.append(combined_word)


        i+=1
        
        
    all_skills=primary_skills_present+primary_skills_present_bi+primary_skills_present_tri+primary_skills_present_tetra
            
        
    return all_skills



class resume:
    def __init__(self,p,ps,ss,e):
        self.path=p
        self.primarySkill=ps
        self.secSkill=ss
        self.exp=e
        
    def getPath(self):
        return self.path
    def getPrimarySkill(self):
        return self.primarySkill
    def getSecSkill(self):
        return self.secSkill
    def getExp(self):
        return self.exp
    def getName(self):
        split_path=self.path.split("/")
        return split_path[-1]
        
    def getVector(self):
        vector=[]
        for s in parent_list:
            if s in self.primarySkill:
                vector.append(1)
            else:
                vector.append(0)
        for s in secondary_list:
            if s in self.primarySkill:
                vector.append(1)
            else:
                vector.append(0)
        vector.append(self.exp)
#         vector=np.array(vector)
        return vector
            

def getLabelName(n):
    if n==1:
        return 'Senior Developer'
    elif n==0:
        return 'Junior Developer'

obj_list=[]
document_list = []
slider_value=50

    
def append_obj(document_path,k):
    final_words=shortenDoc(document_path)
    primary_skills_present=primarySkillMatch(final_words)
    secondary_skills_present=secSkillMatch(final_words)
    exp=experience(final_words)
    obj=resume(document_path,primary_skills_present,secondary_skills_present,exp)
    obj_list.append(obj)
    if k+1<len(document_list):
        append_obj(document_list[k+1],k+1)
    return
        
def allToExcel(document_directory):
    k=0
    if(len(document_directory)==0):
        return
    
    for path, subdirs, files in os.walk(document_directory): 
        for name in files:
            if os.path.splitext(os.path.join(path, name))[1] == ".docx":
                document_list.append(os.path.join(path, name))
    
    if(len(document_list)==0):
        return
    
    append_obj(document_list[k],k)

    
    
    final_vector=[]
    for i in range (0,len(document_list)):
        final_vector.append(obj_list[i].getVector())

    final_vector=np.array(final_vector)

    from sklearn.ensemble import RandomForestClassifier
    from sklearn.model_selection import train_test_split
    from sklearn.metrics import accuracy_score
    from sklearn.neighbors import KNeighborsClassifier
    import pickle


    filename = 'finalized_model.sav'
    loaded_model = pickle.load(open(filename, 'rb'))
    label_list = loaded_model.predict(final_vector)
    # result = loaded_model.score(data_test, label_test)
    # print(result)

    try:
        with open('allResumeData.csv','w',) as f:
            thewriter=csv.writer(f)
            thewriter.writerow(['Sr. No.','File Name','Primary Skills','Secondary Skills','Experience','Eligibility'])

            for i in range (0,len(document_list)):
                thewriter.writerow([i+1,obj_list[i].getName(),obj_list[i].getPrimarySkill(),obj_list[i].getSecSkill(),obj_list[i].getExp(),getLabelName(label_list[i])])
        #         print("resume ",i+1,":-\nfile name: ",obj_list[i].getName(),"\nskills: ",obj_list[i].getSkill(),"\nExperience: ",obj_list[i].getExp(),"\n\n")


        print("All Resume's data has been exported to allResumeData.csv")
    except:
        print("File is in use", "The excel file is in use by another program, kindly close it and try again by clicking compile data.")
            

        
def allToExcelGUI():
    threading.Thread(target=allToExcel).start()
    
    
    
    
def filterToExcel():
    user_description=input('enter description ')
    final_words=shortenText(user_description)
    job_description=primarySkillMatch(final_words)

    with open('Recommended.csv','w',) as f:
        thewriter=csv.writer(f)
        thewriter.writerow(['Sr. No.','File Name','Primary Skills','Secondary Skills','Experience','Eligibility'])
        sr=1

        for i in range (0,len(document_list)):
            skill_list=obj_list[i].getPrimarySkill()
            if listMatch(skill_list,job_description):
                thewriter.writerow([sr,obj_list[i].getName(),obj_list[i].getPrimarySkill(),obj_list[i].getSecSkill(),obj_list[i].getExp(),getLabelName(label_list[i])])
                sr+=1

def recommend(jd,slider_value):
#     print (slider_value)
    import ast
    try:
        user_description=jd
        if user_description=='':
            print("No query entered", "Please enter a query into the query box to get recommendations")
            return None
        
    except:
        user_description=input('enter job description: ')
    final_words=shortenText(user_description)
    job_description=primarySkillMatch(final_words)+secSkillMatch(final_words)

    try:
        csv_file = csv.reader(open('allResumeData.csv', "r"))
    except:
        print("Unsuccesful!", "Resume data has not been already made, kindly make resume data first by clicking Compile Data button")
        return
    
    
    with open('Recommended.csv','w',) as f:
        thewriter=csv.writer(f)
        thewriter.writerow(['Sr. No.','File Name','Primary Skills','Secondary Skills','Experience','Eligibility'])

        
        sr=0
        next(csv_file)

        for data in csv_file:
            if(data.__len__()==0):
                continue
            else:
                primary_skill_list=ast.literal_eval(data[2])
                sec_skill_list=ast.literal_eval(data[3])
                skill_list=primary_skill_list+sec_skill_list
                if listMatch(skill_list,job_description,slider_value):
                    thewriter.writerow([sr+1, data[1],data[2],data[3],data[4],data[5]])
                    sr+=1
        
    try:
        print("Successfully Completed", "Recommended resumes's data has been successfully exported to recommended.csv")
    except:
        print("Recommended resumes's data has been exported to recommended.csv")

        
def recommendInGUI():
    
    user_description='java'
    if user_description=='':
        print("No query entered", "Please enter a query into the query box to get recommendations")
        return None
    

    def inner():
        try:
            recommend(90)
        except:
            print("Unsuccesful!", "The Excel file is open in another program, kindly close it and try again")
            return


    inner()
#         top.mainloop()


def commandLine():
    while(True):
        x=int(input("What do you want to do? \n1. Make Resume data\n2. Recommend Resumes\n3. Exit\n"))
        if x==3:
            break
        elif x==1:
            allToExcel()
            
        elif x==2:
            try:
                recommend()
            except:
                print("Resume data has not been already made, kindly make resume data first")
                



import threading

